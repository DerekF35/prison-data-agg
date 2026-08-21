#!/usr/bin/env python3
"""
Generates high-quality Word (.docx) and PDF (.pdf) documents for the
US Correctional Facilities Master Dataset Methodology Report.
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import subprocess

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
DOCX_OUT = os.path.join(OUTPUT_DIR, "US_Correctional_Facilities_Methodology_Report.docx")
PDF_OUT = os.path.join(OUTPUT_DIR, "US_Correctional_Facilities_Methodology_Report.pdf")

doc = docx.Document()

# Set standard 1-inch margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    # Add page number in footer
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = f_p.add_run("US Correctional Facilities Master Dataset • Page ")
    f_run.font.name = "Calibri"
    f_run.font.size = Pt(9)
    f_run.font.color.rgb = RGBColor(128, 128, 128)

# Colors
PRIMARY_NAVY = RGBColor(31, 78, 120)     # #1F4E78
SECONDARY_BLUE = RGBColor(46, 117, 182)  # #2E75B6
DARK_GRAY = RGBColor(51, 51, 51)         # #333333
LIGHT_BG_HEX = "F2F5F9"
NAVY_BG_HEX = "1F4E78"
BORDER_HEX = "D9D9D9"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_HEX}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_HEX}"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_HEX}"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

# --- Title Header ---
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(4)
title_run = title_p.add_run("US Correctional Facilities Master Dataset")
title_run.font.name = "Calibri"
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = PRIMARY_NAVY

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(18)
sub_run = sub_p.add_run("Comprehensive Technical Documentation & Methodology Report")
sub_run.font.name = "Calibri"
sub_run.font.size = Pt(13)
sub_run.font.italic = True
sub_run.font.color.rgb = SECONDARY_BLUE

# --- Callout Box: Executive Overview ---
overview_table = doc.add_table(rows=1, cols=1)
overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
overview_table.autofit = False
overview_cell = overview_table.cell(0, 0)
overview_cell.width = Inches(6.5)
set_cell_background(overview_cell, "F2F5F9")
set_cell_margins(overview_cell, top=140, bottom=140, left=200, right=200)

box_p = overview_cell.paragraphs[0]
box_p.paragraph_format.space_before = Pt(0)
box_p.paragraph_format.space_after = Pt(4)
b_title = box_p.add_run("EXECUTIVE SUMMARY & KEY DATASET METRICS\n")
b_title.font.bold = True
b_title.font.size = Pt(11)
b_title.font.color.rgb = PRIMARY_NAVY

metrics_text = (
    "• Total Master Facilities: 6,768 unique institutions nationwide\n"
    "• Geographic Coverage: 55 States & Territories (All 50 states + DC, PR, GU, VI, MP)\n"
    "• GPS Coordinates Completeness: 100.0% (6,768 / 6,768 facilities validated in WGS84)\n"
    "• Physical Street Addresses: 100.0% (6,765 verified street addresses)\n"
    "• Direct Contact Phone Numbers: 6,229 facilities\n"
    "• Total Reported Design Bed Capacity: 2,411,708 beds\n"
    "• Total Reported Inmate Population: 2,069,547 inmates"
)
b_body = box_p.add_run(metrics_text)
b_body.font.size = Pt(10)
b_body.font.color.rgb = DARK_GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Helper: Section Header
def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_NAVY
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = SECONDARY_BLUE
    return p

def add_body_p(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY
    return p

# Section 1
add_heading_1("1. Research Context & Data Provenance")
add_body_p(
    "The United States correctional infrastructure is decentralized across federal, state, county, "
    "municipal, tribal, and private jurisdictions. No single governmental agency maintains a real-time, unified master registry "
    "of all facilities. To build this comprehensive national database, data was acquired and synthesized from two authoritative federal sources:"
)
add_body_p(
    "1. Homeland Infrastructure Foundation-Level Data (HIFLD) – Prison Facilities Layer (DHS / FEMA / Oak Ridge National Laboratory):\n"
    "   Provides critical infrastructure geospatial baseline data for secure detention facilities across all 50 states, DC, and territories."
)
add_body_p(
    "2. Federal Bureau of Prisons (BOP) Official Public Institution Directory (U.S. Department of Justice):\n"
    "   Provides authoritative live operational records for all Federal Correctional Complexes (FCC), US Penitentiaries (USP), "
    "Federal Correctional Institutions (FCI), Federal Prison Camps (FPC), Detention Centers (FDC/MDC), and Reentry Offices (RRM)."
)

# Section 2: Table of Jurisdictions
add_heading_1("2. Jurisdictional Breakdown & Summary")
add_body_p("The dataset captures 6,768 facilities across six jurisdictional authorities:")

jur_table_data = [
    ["Jurisdiction", "Facility Count", "Design Bed Capacity", "Reported Population", "Primary Facility Classifications"],
    ["County / Local", "3,924", "777,361", "608,074", "County Jails, Adult Detention Centers, Juvenile Facilities"],
    ["State", "2,347", "1,364,812", "1,223,733", "State Prisons, Correctional Institutions, Re-entry Centers"],
    ["Federal", "253", "186,134", "165,862", "Federal Bureau of Prisons (USP, FCI, FPC, FDC, MDC, FMC, RRM)"],
    ["Municipal / Local", "182", "38,490", "32,841", "City Jails, Municipal Holding Facilities"],
    ["Multi-Jurisdiction", "35", "16,985", "15,221", "Regional Jail Authorities, Joint County Compacts"],
    ["Private / Contract", "27", "27,926", "23,816", "Contracted Detention Facilities"],
    ["Total", "6,768", "2,411,708", "2,069,547", "Nationwide Master Total"]
]

jur_table = doc.add_table(rows=len(jur_table_data), cols=5)
jur_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(jur_table)

col_widths = [Inches(1.3), Inches(0.9), Inches(1.1), Inches(1.1), Inches(2.1)]

for r_idx, row_data in enumerate(jur_table_data):
    row = jur_table.rows[r_idx]
    is_header = (r_idx == 0)
    is_total = (r_idx == len(jur_table_data) - 1)
    for c_idx, cell_value in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = col_widths[c_idx]
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        # Alignment
        if c_idx in [1, 2, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        run = p.add_run(cell_value)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        
        if is_header:
            set_cell_background(cell, NAVY_BG_HEX)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_total:
            set_cell_background(cell, "EAECEF")
            run.font.bold = True
        elif r_idx % 2 == 1:
            set_cell_background(cell, LIGHT_BG_HEX)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Section 3
add_heading_1("3. Ingestion, Cleaning & Entity Resolution Methodology")

add_heading_2("A. Ingestion & API Pagination")
add_body_p(
    "Data was extracted from the DHS HIFLD ArcGIS FeatureServer using programmatic HTTP pagination with parameters "
    "resultOffset and resultRecordCount=2000 in WGS84 spatial reference (outSR=4326). This bypassed server transfer limits "
    "and ensured 100% feature capture. Raw responses were cached deterministically to guarantee reproducibility."
)

add_heading_2("B. Deduplication of Spatial Multi-Part Features")
add_body_p(
    "ArcGIS layers frequently represent multi-building complexes or split campus parcels as separate polygon geometries, "
    "resulting in duplicate attribute entries. The raw HIFLD layer contained 10,738 entries representing exactly 6,737 unique physical facilities. "
    "The pipeline indexed facilities strictly by unique FACILITYID, retaining the primary feature with verified coordinates."
)

add_heading_2("C. Collision-Free Entity Matching")
add_body_p(
    "To integrate live Federal BOP directory records with the HIFLD baseline without corrupting local municipal entities, "
    "a strict entity resolution algorithm was enforced:\n"
    "• Federal BOP codes (e.g. BOP-MAR, BOP-THA) and exact institution titles were matched directly.\n"
    "• Substring matching was strictly restricted to facilities classified as Federal or containing federal institution markers (USP, FCI, FDC, MDC, FMC).\n"
    "• This prevented false collisions where short federal institution names would otherwise overwrite local county jails in the same municipality (e.g. USP Marion vs Marion County Jail in Illinois)."
)

add_heading_2("D. Data Sanitization & Formatting Rules")
add_body_p(
    "• Sentinel Values: Purged legacy placeholder strings ('-999', '9999', 'NOT AVAILABLE', '-1--1') into clean null values.\n"
    "• Title Casing & Acronym Preservation: Standardized names to Title Case while whitelisting uppercase federal acronyms (USP, FCI, ADX, FDC, MDC, FMC, BOP, DOC, USMS, ICE).\n"
    "• Postal ZIP & FIPS Formatting: Preserved leading zeroes for East Coast states and territories (e.g. '01862' for MA, '00921' for PR). Corrected upstream Pickens County Alabama FIPS typo (10107 -> 01107).\n"
    "• Discrete Integer Representation: Capacities and populations are serialized as clean nullable integers without float decimals."
)

# Section 4: Top States Table
add_heading_1("4. Geographic Distribution (Top 10 States)")

state_table_data = [
    ["Rank", "State / Territory", "Facility Count", "Reported Capacity", "State DOC", "County Jails", "Federal"],
    ["1", "Texas (TX)", "555", "311,736", "136", "361", "36"],
    ["2", "Florida (FL)", "417", "178,382", "286", "103", "14"],
    ["3", "California (CA)", "413", "215,306", "101", "267", "26"],
    ["4", "Georgia (GA)", "324", "116,652", "88", "199", "12"],
    ["5", "Ohio (OH)", "237", "76,408", "40", "154", "7"],
    ["6", "New York (NY)", "232", "89,792", "68", "99", "12"],
    ["7", "North Carolina (NC)", "227", "58,671", "67", "131", "9"],
    ["8", "Missouri (MO)", "199", "50,460", "31", "138", "7"],
    ["9", "Virginia (VA)", "192", "57,230", "48", "101", "13"],
    ["10", "Illinois (IL)", "187", "68,284", "45", "109", "8"]
]

state_table = doc.add_table(rows=len(state_table_data), cols=7)
state_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(state_table)

col_w_state = [Inches(0.6), Inches(1.4), Inches(1.0), Inches(1.2), Inches(0.8), Inches(0.8), Inches(0.7)]

for r_idx, row_data in enumerate(state_table_data):
    row = state_table.rows[r_idx]
    is_header = (r_idx == 0)
    for c_idx, cell_value in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = col_w_state[c_idx]
        set_cell_margins(cell, top=70, bottom=70, left=80, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        if c_idx in [0, 2, 3, 4, 5, 6]:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        run = p.add_run(cell_value)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        
        if is_header:
            set_cell_background(cell, NAVY_BG_HEX)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif r_idx % 2 == 1:
            set_cell_background(cell, LIGHT_BG_HEX)

doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Section 5: Data Dictionary Table
add_heading_1("5. Schema & Field Reference")

dict_table_data = [
    ["Column Name", "Type", "Example", "Description"],
    ["facility_id", "String", "10002798", "Primary unique alphanumeric facility identifier (HIFLD ID / BOP Code)."],
    ["facility_name", "String", "Midland Co Central Det", "Official standardized title-cased facility name."],
    ["jurisdiction", "String", "County / Local", "Authority level: Federal, State, County/Local, Municipal, Private."],
    ["facility_type", "String", "County / Local Jail", "Operational classification (State/Federal Prison, Jail, Juvenile, etc.)."],
    ["security_level", "String", "Maximum", "Security level (Maximum, Close, Medium, Minimum, Juvenile, Admin)."],
    ["operational_status", "String", "Open", "Operational status: Open, Closed, Not Available."],
    ["street_address", "String", "400 S Main St", "Physical street address of the facility."],
    ["city", "String", "Midland", "City where the facility is located."],
    ["state", "String", "TX", "Two-letter US postal state/territory abbreviation (55 total)."],
    ["zip_code", "String", "79701", "5-digit or 9-digit postal ZIP code (leading zeroes preserved)."],
    ["county", "String", "Midland", "County or parish name."],
    ["county_fips", "String", "48329", "5-digit Federal Information Processing Standard county code."],
    ["phone_number", "String", "(432) 688-4745", "Standardized 10-digit telephone contact number."],
    ["website", "String", "https://...", "Official facility or department portal URL."],
    ["latitude", "Float", "31.993959", "WGS84 Decimal Degrees Latitude (North)."],
    ["longitude", "Float", "-102.075419", "WGS84 Decimal Degrees Longitude (West)."],
    ["design_capacity", "Integer", "498", "Official rated or design bed capacity."],
    ["population", "Integer", "438", "Reported inmate population count."],
    ["gender", "String", "Male / Female", "Inmate gender housing designation."],
    ["data_source", "String", "DHS HIFLD", "Origin agency of baseline record."]
]

dict_table = doc.add_table(rows=len(dict_table_data), cols=4)
dict_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(dict_table)

col_w_dict = [Inches(1.4), Inches(0.7), Inches(1.3), Inches(3.1)]

for r_idx, row_data in enumerate(dict_table_data):
    row = dict_table.rows[r_idx]
    is_header = (r_idx == 0)
    for c_idx, cell_value in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = col_w_dict[c_idx]
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        run = p.add_run(cell_value)
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        
        if is_header:
            set_cell_background(cell, "333F48")
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif r_idx % 2 == 1:
            set_cell_background(cell, LIGHT_BG_HEX)

# Save Word document
doc.save(DOCX_OUT)
print(f"[+] Successfully generated Word document: {DOCX_OUT} ({os.path.getsize(DOCX_OUT):,} bytes)")

# Convert to PDF via LibreOffice headless
print("[*] Converting Word document to PDF using LibreOffice headless...")
try:
    cmd = [
        "libreoffice", "--headless", "--convert-to", "pdf",
        DOCX_OUT, "--outdir", OUTPUT_DIR
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if os.path.exists(PDF_OUT):
        print(f"[+] Successfully generated PDF document: {PDF_OUT} ({os.path.getsize(PDF_OUT):,} bytes)")
    else:
        print(f"[-] PDF conversion completed but file not found at {PDF_OUT}")
except Exception as e:
    print(f"[-] LibreOffice PDF conversion error: {e}")
